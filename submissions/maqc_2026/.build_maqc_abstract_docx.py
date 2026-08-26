from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("MAQC_2026_SpaceBioBench_Abstract_Kim_Mason.docx")

TITLE = (
    "From Reproducible Pipelines to Reproducible Claims: "
    "An Audit of Model Selection and Aggregation in Spaceflight Omics"
)

PARAGRAPHS = [
    (
        "Technical reproducibility is necessary but not sufficient for a reproducible scientific claim. "
        "Spaceflight omics is a demanding test case because small cohorts and mission-specific hardware, "
        "protocols, age, tissue handling, and processing are entangled with biological variation. We "
        "examined how model selection and score aggregation affect conclusions about model performance "
        "under cross-study shift."
    ),
    (
        "SpaceBio-Bench contains 549 mouse bulk RNA-seq profiles from six tissues organized into 22 "
        "leave-one-mission-out task folds. We audited a fixed PCA-logistic regression pipeline and two "
        "widely used first-generation single-cell-pretrained models, scGPT and Mouse-Geneformer, selected "
        "because their fold-level training histories were available across the audited benchmark surface. "
        "This was a retrospective audit rather than a comprehensive comparison of current models. It "
        "separated two choices that had been combined in the initial summary: selecting an epoch using the "
        "held-out mission and defining performance as either the mean of mission-level AUROCs or a pooled "
        "out-of-fold AUROC. Result coverage was retained explicitly: 21 folds for scGPT and 22 for "
        "Mouse-Geneformer."
    ),
    (
        "The initial summary selected each foundation model's best held-out test epoch. In a fixed epoch-10 "
        "sensitivity analysis, the six-tissue macro AUROC decreased from 0.666 to 0.599 for scGPT and from "
        "0.476 to 0.458 for Mouse-Geneformer; a fixed PCA-logistic regression pipeline achieved 0.730. "
        "Within this audited benchmark surface, neither tested model outperformed the conventional pipeline, "
        "and the estimated gaps increased under the fixed-epoch analysis. Aggregation also changed the "
        "reported effect. For thymus, the PCA-logistic regression mean mission-level AUROC was 0.923, whereas "
        "its pooled out-of-fold AUROC was 0.631. These values answer different questions, and neither is "
        "interpretable without naming the evaluation unit and aggregation rule."
    ),
    (
        "Mission-held-out evaluation alone therefore did not make the resulting model-comparison claim fully "
        "specified. We propose a minimum reporting profile that records task and fold coverage, train-side "
        "model-selection rules, per-mission and pooled estimands, model and preprocessing versions, and "
        "prediction-level run records. Spaceflight omics provides a public stress test for moving from "
        "rerunnable pipelines to traceable, reproducible claims about AI-enabled science."
    ),
]


def set_run_font(run, *, size, bold=False, italic=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif not value and node is not None:
        p_pr.remove(node)


def set_keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# proposal_centerpiece opening pattern, restrained for a one-page academic abstract.
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_before = Pt(0)
kicker.paragraph_format.space_after = Pt(7)
set_keep_with_next(kicker)
set_run_font(
    kicker.add_run("MAQC SOCIETY ANNUAL MEETING 2026  |  ABSTRACT DRAFT FOR CO-AUTHOR REVIEW"),
    size=9,
    bold=True,
    color="6B7280",
)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
title.paragraph_format.line_spacing = 1.05
set_keep_with_next(title)
set_run_font(title.add_run(TITLE), size=16.5, bold=True, color="17324D")

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_before = Pt(0)
authors.paragraph_format.space_after = Pt(2)
set_keep_with_next(authors)
set_run_font(authors.add_run("JangKeun Kim, Christopher Mason"), size=11.5, bold=True)

affiliation = doc.add_paragraph()
affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
affiliation.paragraph_format.space_before = Pt(0)
affiliation.paragraph_format.space_after = Pt(16)
set_keep_with_next(affiliation)
set_run_font(
    affiliation.add_run("Weill Cornell Medicine, New York, NY, USA"),
    size=10.5,
    italic=True,
    color="4B5563",
)

heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
heading.paragraph_format.space_before = Pt(0)
heading.paragraph_format.space_after = Pt(7)
set_keep_with_next(heading)
set_run_font(heading.add_run("Abstract"), size=13, bold=True, color="17324D")

for text in PARAGRAPHS:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.10
    set_keep_lines(paragraph)
    set_run_font(paragraph.add_run(text), size=10.7)

doc.core_properties.title = TITLE
doc.core_properties.subject = "MAQC Society Annual Meeting 2026 abstract draft"
doc.core_properties.author = "JangKeun Kim; Christopher Mason"
doc.core_properties.keywords = "MAQC; spaceflight omics; reproducibility; foundation models; SpaceBio-Bench"

doc.save(OUTPUT)
print(OUTPUT)
